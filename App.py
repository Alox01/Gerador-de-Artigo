from flask import Flask, request, send_from_directory, render_template
import os
import google.generativeai as genai
import re
from servico_banco import salvar_trabalho
from servico_banco import listar_trabalhos
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt, RGBColor
from docx.shared import Cm
from reportlab.lib.pagesizes import A4
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import cm
from reportlab.lib.enums import TA_JUSTIFY
from dotenv import load_dotenv
from reportlab.platypus import KeepTogether
from reportlab.pdfgen import canvas
from flask import render_template 
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
import webbrowser
import threading

# Carregar a chave de API do arquivo .env
load_dotenv()

app = Flask(__name__)

# Configuração da API Gemini
api_key = os.getenv("API_KEY")
genai.configure(api_key=api_key)
modelo = genai.GenerativeModel(model_name="gemini-2.0-flash")

# Diretório para salvar os arquivos gerados
OUTPUT_DIR = "output_files"
if not os.path.exists(OUTPUT_DIR):
    os.makedirs(OUTPUT_DIR)

# Função para formatar o título com a primeira letra de cada palavra em maiúscula
def formatar_titulo(titulo):
    return titulo.title()

# Função para remover os asteriscos de qualquer parte do texto
def remover_asteriscos(texto):
    return re.sub(r"\*+", "", texto)

def adicionar_paginacao(canvas_obj, doc_obj):
    canvas_obj.setFont("Times-Roman", 12)
    pagina = canvas_obj.getPageNumber()
    largura, altura = A4
    canvas_obj.drawRightString(largura - 2 * cm, altura - 1.5 * cm, str(pagina))
    
def adicionar_num_pagina_word(doc):
    section = doc.sections[0]
    header = section.header
    par = header.paragraphs[0] if header.paragraphs else header.add_paragraph()
    par.alignment = WD_ALIGN_PARAGRAPH.RIGHT

    # Inserir campo de número de página
    run = par.add_run()
    fldChar1 = OxmlElement('w:fldChar')
    fldChar1.set(qn('w:fldCharType'), 'begin')
    instrText = OxmlElement('w:instrText')
    instrText.text = 'PAGE'
    fldChar2 = OxmlElement('w:fldChar')
    fldChar2.set(qn('w:fldCharType'), 'end')
    run._r.append(fldChar1)
    run._r.append(instrText)
    run._r.append(fldChar2)
    run.font.name = 'Times New Roman'
    run.font.size = Pt(12)

# Função para formatar parágrafos com recuo de 1,25 cm e espaçamento entre parágrafos de 1,5
def formatar_paragrafos(doc):
    for paragrafo in doc.paragraphs:
        if not paragrafo.style.name.startswith('Heading'):
            paragrafo.paragraph_format.left_indent = Cm(1.25)
            paragrafo.paragraph_format.line_spacing = 1.5
        else:
            paragrafo.paragraph_format.left_indent = Cm(0)
            


# Função para gerar o artigo
def gerar_artigo_abnt(titulo, tema, autor):
    titulo = formatar_titulo(titulo)

    prompt = f"""
    Gere um artigo acadêmico completo e bem estruturado sobre **"{tema}"** com o título **"{titulo}"**, seguindo rigorosamente as normas da ABNT. O artigo deve conter as seguintes seções obrigatórias, com seus respectivos conteúdos e tamanhos mínimos:

    **1. Título**
    - Deve aparecer centralizado no início.

    **2. Resumo (em português)**
    - Um parágrafo entre 150 a 250 palavras que sintetize os principais pontos do artigo.
    
    **3. Palavras-chave (em português)**
    - De 3 a 5 palavras separadas por ponto e vírgula (;).

    **4. Abstract (em inglês)**
    - Um parágrafo com a tradução do resumo, entre 150 a 250 palavras.
    - Sintetize os principais pontos do artigo em inglês.
    
    **5. Keywords (em inglês)**
    - Tradução das palavras-chave, entre 3 e 5 termos separados por ponto e vírgula (;).

    **6. Introdução**
    - Apresente o tema, justificativa, problema e objetivo da pesquisa.
    - Mínimo de 200 palavras.

    **7. Revisão de Literatura**
    - Discorra sobre conceitos teóricos importantes sobre o tema.
    - Utilize ao menos 2 citações no estilo ABNT: (SOBRENOME, ano, p.xx).
    - Mínimo de 300 palavras.

    **8. Metodologia**
    - Descreva os métodos e procedimentos adotados para desenvolver o trabalho.
    - Pode incluir abordagem qualitativa/quantitativa, revisão bibliográfica, etc.
    - Mínimo de 200 palavras.

    **9. Resultados e Discussão**
    - Apresente os principais resultados esperados ou obtidos.
    - Relacione com a literatura citada.
    - Mínimo de 300 palavras.

    **10. Conclusão**
    - Retome os objetivos, destaque as contribuições e proponha trabalhos futuros.
    - Mínimo de 150 palavras.

    **11. Referências**
    - Liste pelo menos **3 referências no formato ABNT.**
    - Exemplo: SOBRENOME, Nome. *Título do Livro ou Artigo*. Local: Editora, Ano.
    """
    resposta = modelo.generate_content(prompt)
    return resposta.text

# Função para salvar o texto no PDF com formatação correta
def salvar_em_pdf(titulo, texto, autor):
    nome_arquivo = os.path.join(OUTPUT_DIR, titulo.replace(" ", "_") + ".pdf")
    doc = SimpleDocTemplate(nome_arquivo, pagesize=A4,
                            rightMargin=2*cm, leftMargin=3*cm,
                            topMargin=3*cm, bottomMargin=2*cm)

    styles = getSampleStyleSheet()
    styles.add(ParagraphStyle(name='Titulo', fontSize=12, fontName='Times-Roman', alignment=1,
                              spaceAfter=12, spaceBefore=12, leading=14, allowHTML=True))
    styles.add(ParagraphStyle(name='Secao', fontSize=12, fontName='Times-Roman', alignment=0,
                              spaceAfter=12, spaceBefore=12, leading=14, firstLineIndent=0, allowHTML=True))
    styles.add(ParagraphStyle(name='Texto', fontSize=12, fontName='Times-Roman', alignment=TA_JUSTIFY,
                              firstLineIndent=1.25 * cm, spaceAfter=10, leading=18, allowHTML=True))
    styles.add(ParagraphStyle(name='TextoSemRecuo', fontSize=12, fontName='Times-Roman', alignment=TA_JUSTIFY,
                              firstLineIndent=0, spaceAfter=10, leading=18, allowHTML=True))

    elementos = []

    # Título do trabalho
    titulo_upper = titulo.upper()
    elementos.append(Paragraph(f"<b>{titulo_upper}</b>", styles['Titulo']))
    elementos.append(Spacer(1, 0.6*cm))

    # Autor
    autor_sem_quebra = autor.replace(" ", "\u00A0")
    autor_para_pdf = Paragraph(
        autor_sem_quebra,
        ParagraphStyle(
            'AutorDireita',
            parent=styles['TextoSemRecuo'],
            alignment=2,
            spaceAfter=12,
            fontName='Times-Roman',
            fontSize=12,
        )
    )
    elementos.append(autor_para_pdf)
    elementos.append(Spacer(1, 1.2 * cm))

    # Regex para seções
    padroes = {
        "Resumo": re.compile(r"^\s*\d{0,2}\.?\s*resumo\s*:?\s*(.*)$", re.IGNORECASE),
        "Palavras-chave": re.compile(r"^\s*\d{0,2}\.?\s*palavras-chave\s*:?\s*(.*)$", re.IGNORECASE),
        "Abstract": re.compile(r"^\s*\d{0,2}\.?\s*abstract\s*:?\s*(.*)$", re.IGNORECASE),
        "Keywords": re.compile(r"^\s*\d{0,2}\.?\s*keywords\s*:?\s*(.*)$", re.IGNORECASE),
        "Introdução": re.compile(r"^\s*\d{0,2}\.?\s*introdução\s*:?\s*(.*)$", re.IGNORECASE),
        "Revisão de Literatura": re.compile(r"^\s*\d{0,2}\.?\s*revis[aã]o de literatura\s*:?\s*(.*)$", re.IGNORECASE),
        "Metodologia": re.compile(r"^\s*\d{0,2}\.?\s*metodologia\s*:?\s*(.*)$", re.IGNORECASE),
        "Resultados e Discussão": re.compile(r"^\s*\d{0,2}\.?\s*resultados e discussão\s*:?\s*(.*)$", re.IGNORECASE),
        "Conclusão": re.compile(r"^\s*\d{0,2}\.?\s*conclus[aã]o\s*:?\s*(.*)$", re.IGNORECASE),
        "Referências": re.compile(r"^\s*\d{0,2}\.?\s*refer[eê]ncias\s*:?\s*(.*)$", re.IGNORECASE)
    }

    conteudo = {}
    atual = None
    aguardando_conteudo_direto = False
    ultima_secao = None

    for linha in texto.splitlines():
        linha_limpa = remover_asteriscos(linha.strip())
        if not linha_limpa:
            continue

        mudou_secao = False
        for nome_secao, padrao in padroes.items():
            match = padrao.match(linha_limpa)
            if match:
                atual = nome_secao
                if atual not in conteudo:
                    conteudo[atual] = []

                inline = match.group(1).strip()
                if inline:
                    conteudo[atual].append(inline)
                    aguardando_conteudo_direto = False
                else:
                    aguardando_conteudo_direto = True
                    ultima_secao = atual
                mudou_secao = True
                break

        if not mudou_secao:
            if aguardando_conteudo_direto and ultima_secao:
                conteudo[ultima_secao].append(linha_limpa)
            elif atual:
                conteudo[atual].append(linha_limpa)

    secoes_ordenadas = [
        ("1", "Introdução"),
        ("2", "Revisão de Literatura"),
        ("3", "Metodologia"),
        ("4", "Resultados e Discussão"),
        ("5", "Conclusão")
    ]
    extras = ["Resumo", "Palavras-chave", "Abstract", "Keywords"]

    for entrada in extras + [nome for _, nome in secoes_ordenadas]:
        if entrada in extras:
            titulo = None
        else:
            numero = next((num for num, nome in secoes_ordenadas if nome == entrada), "")
            titulo = f"<b>{numero} {entrada.upper()}</b>"
            elementos.append(Paragraph(titulo, styles['Secao']))

        if entrada in conteudo and conteudo[entrada]:
            for i, paragrafo in enumerate(conteudo[entrada]):
                if entrada in extras and i == 0:
                    texto_completo = f"<b>{entrada.upper()}:</b> {paragrafo}"
                    estilo = styles['TextoSemRecuo']
                else:
                    texto_completo = paragrafo
                    estilo = styles['Texto']
                elementos.append(Paragraph(texto_completo, estilo))
                elementos.append(Spacer(1, 0.4 * cm))
        else:
            elementos.append(Paragraph("Conteúdo não disponível.", styles['Texto']))

    if "Referências" in conteudo:
        elementos.append(Paragraph("<b>REFERÊNCIAS</b>", styles['Secao']))
        for paragrafo in conteudo["Referências"]:
            elementos.append(Paragraph(paragrafo, styles['TextoSemRecuo']))
            elementos.append(Spacer(1, 0.4 * cm))

    doc.build(elementos, onFirstPage=adicionar_paginacao, onLaterPages=adicionar_paginacao)
    return nome_arquivo

# Função para salvar no DOCX com formatação correta
def salvar_em_docx(titulo, texto, autor):
    doc = Document()
    titulo_original = titulo  # Para nome do arquivo

    # Título principal
    titulo_paragrafo = doc.add_paragraph()
    titulo_paragrafo.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_titulo = titulo_paragrafo.add_run(titulo.upper())
    run_titulo.font.name = "Times New Roman"
    run_titulo.font.size = Pt(12)
    run_titulo.bold = True
    run_titulo.font.color.rgb = RGBColor(0, 0, 0)

    # Autor
    par_autor = doc.add_paragraph(autor)
    par_autor.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    par_autor.paragraph_format.space_after = Pt(20)
    par_autor.paragraph_format.space_before = Pt(24)
    run_autor = par_autor.runs[0]
    run_autor.font.name = "Times New Roman"
    run_autor.font.size = Pt(12)
    run_autor.font.color.rgb = RGBColor(0, 0, 0)

    # Regex para seções
    padroes = {
        "Resumo": re.compile(r"^\s*\d{0,2}\.?\s*resumo\s*:?\s*(.*)$", re.IGNORECASE),
        "Palavras-chave": re.compile(r"^\s*\d{0,2}\.?\s*palavras-chave\s*:?\s*(.*)$", re.IGNORECASE),
        "Abstract": re.compile(r"^\s*\d{0,2}\.?\s*abstract\s*:?\s*(.*)$", re.IGNORECASE),
        "Keywords": re.compile(r"^\s*\d{0,2}\.?\s*keywords\s*:?\s*(.*)$", re.IGNORECASE),
        "Introdução": re.compile(r"^\s*\d{0,2}\.?\s*introdução\s*:?\s*(.*)$", re.IGNORECASE),
        "Revisão de Literatura": re.compile(r"^\s*\d{0,2}\.?\s*revis[aã]o de literatura\s*:?\s*(.*)$", re.IGNORECASE),
        "Metodologia": re.compile(r"^\s*\d{0,2}\.?\s*metodologia\s*:?\s*(.*)$", re.IGNORECASE),
        "Resultados e Discussão": re.compile(r"^\s*\d{0,2}\.?\s*resultados e discussão\s*:?\s*(.*)$", re.IGNORECASE),
        "Conclusão": re.compile(r"^\s*\d{0,2}\.?\s*conclus[aã]o\s*:?\s*(.*)$", re.IGNORECASE),
        "Referências": re.compile(r"^\s*\d{0,2}\.?\s*refer[eê]ncias\s*:?\s*(.*)$", re.IGNORECASE)
    }

    conteudo = {}
    atual = None
    aguardando_conteudo_direto = False
    ultima_secao = None

    for linha in texto.splitlines():
        linha_limpa = remover_asteriscos(linha.strip())
        if not linha_limpa:
            continue

        mudou_secao = False
        for nome_secao, padrao in padroes.items():
            match = padrao.match(linha_limpa)
            if match:
                atual = nome_secao
                if atual not in conteudo:
                    conteudo[atual] = []

                inline = match.group(1).strip()
                if inline:
                    conteudo[atual].append(inline)
                    aguardando_conteudo_direto = False
                else:
                    aguardando_conteudo_direto = True
                    ultima_secao = atual
                mudou_secao = True
                break

        if not mudou_secao:
            if aguardando_conteudo_direto and ultima_secao:
                conteudo[ultima_secao].append(linha_limpa)
            elif atual:
                conteudo[atual].append(linha_limpa)

    secoes_ordenadas = [
        ("1", "Introdução"),
        ("2", "Revisão de Literatura"),
        ("3", "Metodologia"),
        ("4", "Resultados e Discussão"),
        ("5", "Conclusão")
    ]
    extras = ["Resumo", "Palavras-chave", "Abstract", "Keywords"]

    # Monta o corpo do artigo
    for entrada in extras + [nome for _, nome in secoes_ordenadas]:
        if entrada in extras:
            titulo_secao = entrada.upper()
        else:
            numero = next(num for num, nome in secoes_ordenadas if nome == entrada)
            titulo_secao = f"{numero} {entrada.upper()}"

        if entrada not in extras:
            heading = doc.add_heading(level=1)
            run_heading = heading.add_run(titulo_secao)
            run_heading.bold = True
            run_heading.font.name = "Times New Roman"
            run_heading.font.size = Pt(12)
            run_heading.font.color.rgb = RGBColor(0, 0, 0)
            heading.paragraph_format.space_after = Pt(12)

        if entrada in conteudo:
            for i, paragrafo in enumerate(conteudo[entrada]):
                if entrada in extras and i == 0:
                    p = doc.add_paragraph()
                    run_titulo = p.add_run(f"{entrada.upper()}: ")
                    run_titulo.bold = True
                    run_titulo.font.name = "Times New Roman"
                    run_titulo.font.size = Pt(12)
                    run_titulo.font.color.rgb = RGBColor(0, 0, 0)

                    run_texto = p.add_run(paragrafo)
                    run_texto.font.name = "Times New Roman"
                    run_texto.font.size = Pt(12)
                    run_texto.font.color.rgb = RGBColor(0, 0, 0)
                    recuo = Cm(0)
                else:
                    p = doc.add_paragraph()
                    run = p.add_run(paragrafo)
                    run.font.name = "Times New Roman"
                    run.font.size = Pt(12)
                    run.font.color.rgb = RGBColor(0, 0, 0)
                    recuo = Cm(0) if entrada in extras else Cm(1.25)

                p.paragraph_format.first_line_indent = recuo
                p.paragraph_format.line_spacing = 1.5
                p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        else:
            doc.add_paragraph("Conteúdo não disponível.")

    # Adiciona Referências no final
    if "Referências" in conteudo:
        heading = doc.add_heading(level=1)
        run_heading = heading.add_run("REFERÊNCIAS")
        run_heading.bold = True
        run_heading.font.name = "Times New Roman"
        run_heading.font.size = Pt(12)
        run_heading.font.color.rgb = RGBColor(0, 0, 0)
        heading.paragraph_format.space_after = Pt(10)

        for paragrafo in conteudo["Referências"]:
            p = doc.add_paragraph()
            run = p.add_run(paragrafo)
            run.font.name = "Times New Roman"
            run.font.size = Pt(12)
            run.font.color.rgb = RGBColor(0, 0, 0)
            p.paragraph_format.first_line_indent = Cm(0)
            p.paragraph_format.line_spacing = 1.5
            p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    nome_limpo = re.sub(r'[^\w\s-]', '', titulo_original).strip()
    nome_final = "_".join(nome_limpo.split())
    nome_arquivo = os.path.join(OUTPUT_DIR, nome_final + ".docx")

    adicionar_num_pagina_word(doc)

    doc.save(nome_arquivo)
    return nome_arquivo

# Rota para a página inicial 
@app.route('/')
def home():
    return render_template('index.html')

# Rota para gerar o trabalho e permitir o download
@app.route('/gerar_trabalho', methods=['POST'])
def gerar_trabalho():
    autor = request.form.get('autor')
    titulo = request.form.get('titulo')
    tema = request.form.get('tema')

    if not titulo or not tema:
        return "Erro: Título e Tema são obrigatórios", 400

    trabalho = gerar_artigo_abnt(titulo, tema, autor)

    # SALVAR NO BANCO
    salvar_trabalho(titulo, tema, autor, trabalho, pdf=True, docx=True)

    # Gerar arquivos
    nome_docx = os.path.basename(salvar_em_docx(titulo, trabalho, autor))
    nome_pdf = os.path.basename(salvar_em_pdf(titulo, trabalho, autor))

    # Mostrar página com os dois botões
    return render_template('download.html', nome_docx=nome_docx, nome_pdf=nome_pdf,
                       preview=trabalho, titulo=titulo, autor=autor)


@app.route('/download/<nome_arquivo>')
def baixar_arquivo(nome_arquivo):
    if '..' in nome_arquivo or nome_arquivo.startswith('/'):
        return "Arquivo inválido", 400
    caminho = os.path.join(OUTPUT_DIR, nome_arquivo)
    if not os.path.isfile(caminho):
        return "Arquivo não encontrado", 404
    return send_from_directory(OUTPUT_DIR, nome_arquivo, as_attachment=True)

@app.route('/trabalhos')
def trabalhos():
    lista = listar_trabalhos()
    return render_template('trabalhos.html', trabalhos=lista)

# Iniciar o servidor Flask
def abrir_navegador():
    webbrowser.open_new("http://127.0.0.1:5000")

@app.route('/baixar_trabalho_editado', methods=['POST'])
def baixar_trabalho_editado():
    texto = request.form.get('texto_editado')
    titulo = request.form.get('titulo')
    autor = request.form.get('autor')
    formato = request.form.get('formato')  # 'docx' ou 'pdf'

    if not texto or not titulo:
        return "Erro: Texto e Título são obrigatórios", 400

    if formato == 'pdf':
        caminho = salvar_em_pdf(titulo, texto, autor)
    else:
        caminho = salvar_em_docx(titulo, texto, autor)

    nome_arquivo = os.path.basename(caminho)
    return send_from_directory(OUTPUT_DIR, nome_arquivo, as_attachment=True)


if __name__ == "__main__":
    threading.Timer(1.0, abrir_navegador).start()
    app.run(debug=True, use_reloader=False)