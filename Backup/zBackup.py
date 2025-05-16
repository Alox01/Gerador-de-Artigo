from flask import Flask, request, send_from_directory, render_template
import os
import google.generativeai as genai
import re
from servico_banco import salvar_trabalho
from servico_banco import listar_trabalhos
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt, RGBColor
from docx.shared import Cm
from reportlab.lib.pagesizes import A4
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import cm
from reportlab.lib.enums import TA_JUSTIFY
from dotenv import load_dotenv

# Carregar a chave de API do arquivo .env
load_dotenv()

app = Flask(__name__)

# Configuração da API Gemini
api_key = os.getenv("API_KEY")
genai.configure(api_key=api_key)
modelo = genai.GenerativeModel(model_name="gemini-2.0-flash")

# Diretório para salvar os arquivos gerados
OUTPUT_DIR = "output_files"
if not os.path.exists(OUTPUT_DIR):
    os.makedirs(OUTPUT_DIR)

# Função para formatar o título com a primeira letra de cada palavra em maiúscula
def formatar_titulo(titulo):
    return titulo.title()

# Função para formatar o autor no formato "SOBRENOME, Nome"
def formatar_autor(nome):
    partes = nome.strip().split()
    if len(partes) < 2:
        return nome  
    sobrenome = partes[-1].upper()
    nome = " ".join(partes[:-1])
    return f"{sobrenome}, {nome}"

# Função para remover os asteriscos de qualquer parte do texto
def remover_asteriscos(texto):
    return re.sub(r"\*+", "", texto)

# Função para formatar parágrafos com recuo de 1,25 cm e espaçamento entre parágrafos de 1,5
def formatar_paragrafos(doc):
    for paragrafo in doc.paragraphs:
        if not paragrafo.style.name.startswith('Heading'):
            paragrafo.paragraph_format.left_indent = Cm(1.25)
            paragrafo.paragraph_format.line_spacing = 1.5
        else:
            paragrafo.paragraph_format.left_indent = Cm(0)


# Função para gerar o artigo
def gerar_artigo_abnt(titulo, tema, autor):
    titulo = formatar_titulo(titulo)

    prompt = f"""
    Gere um artigo acadêmico completo e bem estruturado sobre **"{tema}"** com o título **"{titulo}"**, seguindo rigorosamente as normas da ABNT. O artigo deve conter as seguintes seções obrigatórias, com seus respectivos conteúdos e tamanhos mínimos:

    **1. Título**
    - Deve aparecer centralizado no início.

    **2. Resumo (em português)**
    - Um parágrafo entre 150 a 250 palavras que sintetize os principais pontos do artigo.
    
    **3. Palavras-chave (em português)**
    - De 3 a 5 palavras separadas por ponto e vírgula (;).

    **4. Abstract (em inglês)**
    - Um parágrafo com a tradução do resumo, entre 150 a 250 palavras.
    - Sintetize os principais pontos do artigo em inglês.
    
    **5. Keywords (em inglês)**
    - Tradução das palavras-chave, entre 3 e 5 termos separados por ponto e vírgula (;).

    **6. Introdução**
    - Apresente o tema, justificativa, problema e objetivo da pesquisa.
    - Mínimo de 200 palavras.

    **7. Revisão de Literatura**
    - Discorra sobre conceitos teóricos importantes sobre o tema.
    - Utilize ao menos 2 citações no estilo ABNT: (SOBRENOME, ano, p.xx).
    - Mínimo de 300 palavras.

    **8. Metodologia**
    - Descreva os métodos e procedimentos adotados para desenvolver o trabalho.
    - Pode incluir abordagem qualitativa/quantitativa, revisão bibliográfica, etc.
    - Mínimo de 200 palavras.

    **9. Resultados e Discussão**
    - Apresente os principais resultados esperados ou obtidos.
    - Relacione com a literatura citada.
    - Mínimo de 300 palavras.

    **10. Conclusão**
    - Retome os objetivos, destaque as contribuições e proponha trabalhos futuros.
    - Mínimo de 150 palavras.

    **11. Referências**
    - Liste pelo menos **3 referências no formato ABNT.**
    - Exemplo: SOBRENOME, Nome. *Título do Livro ou Artigo*. Local: Editora, Ano.
    """
    resposta = modelo.generate_content(prompt)
    return resposta.text

# Função para salvar o texto no PDF com formatação correta
def salvar_em_pdf(titulo, texto, autor):
    nome_arquivo = os.path.join(OUTPUT_DIR, titulo.replace(" ", "_") + ".pdf")
    doc = SimpleDocTemplate(nome_arquivo, pagesize=A4,
                            rightMargin=2*cm, leftMargin=2*cm,
                            topMargin=2*cm, bottomMargin=2*cm)

    # Estilos
    styles = getSampleStyleSheet()
    styles.add(ParagraphStyle(name='Titulo', fontSize=16, spaceAfter=12, alignment=1))
    styles.add(ParagraphStyle(name='Secao', fontSize=14, spaceAfter=8, spaceBefore=12, leading=18, alignment=0, firstLineIndent=0, leftIndent=0))
    styles.add(ParagraphStyle(name='Texto', fontSize=12, leading=18, alignment=TA_JUSTIFY, firstLineIndent=1.25 * cm, spaceAfter=10))
    styles.add(ParagraphStyle(name='TextoSemRecuo', fontSize=12, leading=18, alignment=TA_JUSTIFY, firstLineIndent=0, spaceAfter=10))

    elementos = []
    
    from reportlab.platypus import KeepTogether
    
    # Título do trabalho
    titulo_upper = titulo.upper()
    elementos.append(Paragraph(titulo_upper, styles['Titulo']))
    elementos.append(Spacer(1, 0.6*cm))

    autor_formatado = formatar_autor(autor)
    autor_sem_quebra = autor_formatado.replace(" ", "\u00A0")
    autor_para_pdf = Paragraph(autor_sem_quebra, ParagraphStyle(
        'AutorDireita',
        parent=styles['TextoSemRecuo'],
        alignment=2,
        spaceAfter=12,
    ))

    elementos.append(autor_para_pdf)
    elementos.append(Spacer(1, 1.2 * cm))

    # Regex para detectar seções
    padroes = {
    "Resumo": re.compile(r"^\s*\d{0,2}\.?\s*resumo\s*:?\s*(.*)$", re.IGNORECASE),
    "Palavras-chave": re.compile(r"^\s*\d{0,2}\.?\s*palavras-chave\s*:?\s*(.*)$", re.IGNORECASE),
    "Abstract": re.compile(r"^\s*\d{0,2}\.?\s*abstract\s*:?\s*(.*)$", re.IGNORECASE),
    "Keywords": re.compile(r"^\s*\d{0,2}\.?\s*keywords\s*:?\s*(.*)$", re.IGNORECASE),
    "Introdução": re.compile(r"^\s*\d{0,2}\.?\s*introdução\s*:?\s*(.*)$", re.IGNORECASE),
    "Revisão de Literatura": re.compile(r"^\s*\d{0,2}\.?\s*revis[aã]o de literatura\s*:?\s*(.*)$", re.IGNORECASE),
    "Metodologia": re.compile(r"^\s*\d{0,2}\.?\s*metodologia\s*:?\s*(.*)$", re.IGNORECASE),
    "Resultados e Discussão": re.compile(r"^\s*\d{0,2}\.?\s*resultados e discussão\s*:?\s*(.*)$", re.IGNORECASE),
    "Conclusão": re.compile(r"^\s*\d{0,2}\.?\s*conclus[aã]o\s*:?\s*(.*)$", re.IGNORECASE),
    "Referências": re.compile(r"^\s*\d{0,2}\.?\s*refer[eê]ncias\s*:?\s*(.*)$", re.IGNORECASE)
    }

    conteudo = {}
    atual = None
    aguardando_conteudo_direto = False
    ultima_secao = None

    for linha in texto.splitlines():
        linha_limpa = remover_asteriscos(linha.strip())
        if not linha_limpa:
            continue

        mudou_secao = False
        for nome_secao, padrao in padroes.items():
            match = padrao.match(linha_limpa)
            if match:
                atual = nome_secao
                if atual not in conteudo:
                    conteudo[atual] = []

                inline = match.group(1).strip()
                if inline:
                    conteudo[atual].append(inline)
                    aguardando_conteudo_direto = False
                else:
                    aguardando_conteudo_direto = True
                    ultima_secao = atual

                mudou_secao = True
                break

        if not mudou_secao:
            if aguardando_conteudo_direto and ultima_secao:
                if ultima_secao not in conteudo:
                    conteudo[ultima_secao] = []
                conteudo[ultima_secao].append(linha_limpa)
                continue
            if atual:
                if atual not in conteudo:
                    conteudo[atual] = []
                conteudo[atual].append(linha_limpa)


    # Adicionar conteúdo por seção
    for secao in ["Resumo", "Palavras-chave", "Abstract", "Keywords", "Introdução", "Revisão de Literatura", "Metodologia", "Resultados e Discussão", "Conclusão", "Referências"]:
        elementos.append(Paragraph(secao.upper(), styles['Secao']))
        if secao in conteudo and conteudo[secao]:
            for paragrafo in conteudo[secao]:
                estilo = styles['TextoSemRecuo'] if secao in ["Resumo", "Palavras-chave", "Abstract", "Keywords", "Introdução", "Conclusão", "Referências"] else styles['Texto']
                elementos.append(Paragraph(paragrafo, estilo))
                elementos.append(Spacer(1, 0.4 * cm))
        else:
            elementos.append(Paragraph("Conteúdo não disponível.", styles['Texto']))

    print("=== CONTEÚDO FINAL ===")
    for secao, paragrafos in conteudo.items():
        print(f"\n## {secao}")
        for p in paragrafos:
            print("-", p)

    doc.build(elementos)
    print(f"\n✅ PDF salvo como: {nome_arquivo}")
    return nome_arquivo 

# Função para salvar no DOCX com formatação correta
def salvar_em_docx(titulo, texto, autor):
    doc = Document()
    
    # Título principal
    titulo_paragrafo = doc.add_paragraph()
    titulo_paragrafo.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_titulo = titulo_paragrafo.add_run(titulo.upper())
    run_titulo.font.name = "Times New Roman"
    run_titulo.font.size = Pt(16)
    run_titulo.bold = True
    
    # Autor
    par_autor = doc.add_paragraph(formatar_autor(autor))
    par_autor.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    par_autor.paragraph_format.space_after = Pt(20)
    run_autor = par_autor.runs[0]
    run_autor.font.name = "Times New Roman"
    run_autor.font.size = Pt(12)
    run_autor.font.color.rgb = RGBColor(0, 0, 0)

    


    # Regex para detectar seções
    padroes = {
    "Resumo": re.compile(r"^\s*\d{0,2}\.?\s*resumo\s*:?\s*(.*)$", re.IGNORECASE),
    "Palavras-chave": re.compile(r"^\s*\d{0,2}\.?\s*palavras-chave\s*:?\s*(.*)$", re.IGNORECASE),
    "Abstract": re.compile(r"^\s*\d{0,2}\.?\s*abstract\s*:?\s*(.*)$", re.IGNORECASE),
    "Keywords": re.compile(r"^\s*\d{0,2}\.?\s*keywords\s*:?\s*(.*)$", re.IGNORECASE),
    "Introdução": re.compile(r"^\s*\d{0,2}\.?\s*introdução\s*:?\s*(.*)$", re.IGNORECASE),
    "Revisão de Literatura": re.compile(r"^\s*\d{0,2}\.?\s*revis[aã]o de literatura\s*:?\s*(.*)$", re.IGNORECASE),
    "Metodologia": re.compile(r"^\s*\d{0,2}\.?\s*metodologia\s*:?\s*(.*)$", re.IGNORECASE),
    "Resultados e Discussão": re.compile(r"^\s*\d{0,2}\.?\s*resultados e discussão\s*:?\s*(.*)$", re.IGNORECASE),
    "Conclusão": re.compile(r"^\s*\d{0,2}\.?\s*conclus[aã]o\s*:?\s*(.*)$", re.IGNORECASE),
    "Referências": re.compile(r"^\s*\d{0,2}\.?\s*refer[eê]ncias\s*:?\s*(.*)$", re.IGNORECASE)
    }

    conteudo = {}
    atual = None
    aguardando_conteudo_direto = False
    ultima_secao = None

    for linha in texto.splitlines():
        linha_limpa = remover_asteriscos(linha.strip())
        if not linha_limpa:
            continue

        mudou_secao = False
        for nome_secao, padrao in padroes.items():
            match = padrao.match(linha_limpa)
            if match:
                atual = nome_secao
                if atual not in conteudo:
                    conteudo[atual] = []

                inline = match.group(1).strip()
                if inline:
                    conteudo[atual].append(inline)
                    aguardando_conteudo_direto = False
                else:
                    aguardando_conteudo_direto = True
                    ultima_secao = atual

                mudou_secao = True
                break

        if not mudou_secao:
            if aguardando_conteudo_direto and ultima_secao:
                if ultima_secao not in conteudo:
                    conteudo[ultima_secao] = []
                conteudo[ultima_secao].append(linha_limpa)
                continue
            if atual:
                if atual not in conteudo:
                    conteudo[atual] = []
                conteudo[atual].append(linha_limpa)


    # Adicionar conteúdo por seção
    for secao in ["Resumo", "Palavras-chave", "Abstract", "Keywords", "Introdução", "Revisão de Literatura", "Metodologia", "Resultados e Discussão", "Conclusão", "Referências"]:
        heading = doc.add_heading(secao.upper(), level=2)
        heading.paragraph_format.space_after = Pt(10)
        if secao in conteudo and conteudo[secao]:
            for paragrafo in conteudo[secao]:
                p = doc.add_paragraph()
                run = p.add_run(paragrafo)
                run.font.name = "Times New Roman"
                run.font.size = Pt(12)
                run.font.color.rgb = RGBColor(0, 0, 0)
                
                # Formatação direta no parágrafo
                if secao in ["Resumo", "Palavras-chave", "Abstract", "Keywords", "Introdução", "Conclusão", "Referências"]:
                    p.paragraph_format.first_line_indent = Cm(0)
                else:
                    p.paragraph_format.first_line_indent = Cm(1.25)

                
                p.paragraph_format.line_spacing = 1.5
                
                if secao != "Referências":
                    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    print("=== CONTEÚDO FINAL ===")
    for secao, paragrafos in conteudo.items():
        print(f"\n## {secao}")
        for p in paragrafos:
            print("-", p)

    nome_arquivo = os.path.join(OUTPUT_DIR, titulo.replace(" ", "_") + ".docx")
    doc.save(nome_arquivo)
    print(f"\n✅ Documento salvo como: {nome_arquivo}")
    return nome_arquivo

# Rota para a página inicial
from flask import render_template  

@app.route('/')
def home():
    return render_template('index.html')


# Rota para gerar o trabalho e permitir o download
@app.route('/gerar_trabalho', methods=['POST'])
def gerar_trabalho():
    autor = request.form.get('autor')
    titulo = request.form.get('titulo')
    tema = request.form.get('tema')

    if not titulo or not tema:
        return "Erro: Título e Tema são obrigatórios", 400

    trabalho = gerar_artigo_abnt(titulo, tema, autor)

    # SALVAR NO BANCO
    salvar_trabalho(titulo, tema, autor, trabalho, pdf=True, docx=True)

    # Gerar arquivos
    nome_docx = os.path.basename(salvar_em_docx(titulo, trabalho, autor))
    nome_pdf = os.path.basename(salvar_em_pdf(titulo, trabalho, autor))


    # Gerar arquivos
    nome_docx = os.path.basename(salvar_em_docx(titulo, trabalho, autor))
    nome_pdf = os.path.basename(salvar_em_pdf(titulo, trabalho, autor))

    # Mostrar página com os dois botões
    return render_template('download.html', nome_docx=nome_docx, nome_pdf=nome_pdf)

@app.route('/download/<nome_arquivo>')
def baixar_arquivo(nome_arquivo):
    return send_from_directory(OUTPUT_DIR, nome_arquivo, as_attachment=True)

@app.route('/trabalhos')
def trabalhos():
    lista = listar_trabalhos()
    return render_template('trabalhos.html', trabalhos=lista)

# Iniciar o servidor Flask
import webbrowser
import threading

def abrir_navegador():
    webbrowser.open_new("http://127.0.0.1:5000")

if __name__ == "__main__":
    threading.Timer(1.0, abrir_navegador).start()
    app.run(debug=True, use_reloader=False)

