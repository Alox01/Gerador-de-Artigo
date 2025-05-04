# pip install google-generativeai python-docx reportlab python-dotenv Flask

import google.generativeai as genai
import re
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt, RGBColor
from reportlab.lib.pagesizes import A4
from reportlab.pdfgen import canvas
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.platypus import Paragraph, SimpleDocTemplate, Spacer
from reportlab.lib.units import cm
from reportlab.lib.enums import TA_JUSTIFY
from dotenv import load_dotenv
import os

# Carregar a chave de API do arquivo .env
load_dotenv()

# Configuração da API Gemini
api_key = os.getenv("API_KEY")  
genai.configure(api_key=api_key)
modelo = genai.GenerativeModel(model_name="gemini-2.0-flash")

# Função para formatar o título com a primeira letra de cada palavra em maiúscula
def formatar_titulo(titulo):
    return titulo.title()  

# Função para remover os asteriscos de qualquer parte do texto
def remover_asteriscos(texto):
    return re.sub(r"\*+", "", texto) 

# Função para garantir o recuo de 1,25 cm e o espaçamento entre parágrafos de 1,5
def formatar_paragrafos(doc):
    for paragrafo in doc.paragraphs:
        paragrafo.paragraph_format.left_indent = Pt(1.25) 
        paragrafo.paragraph_format.line_spacing = 1.5 

# Função para gerar o artigo
def gerar_artigo_abnt(titulo, tema, nome_autor=None, instituicao=None):
    titulo = formatar_titulo(titulo)

    prompt = f"""
    Gere um artigo acadêmico completo e bem estruturado sobre **"{tema}"** com o título **"{titulo}"**, seguindo rigorosamente as normas da ABNT. O artigo deve conter as seguintes seções obrigatórias, com seus respectivos conteúdos e tamanhos mínimos:

    **1. Título**
    - Deve aparecer centralizado no início.

    **2. Resumo (em português)**
    - Um parágrafo entre 150 a 250 palavras que sintetize os principais pontos do artigo.

    **3. Abstract (em inglês)**
    - Um parágrafo com a tradução do resumo, entre 150 a 250 palavras.
    - Sintetize os principais pontos do artigo em inglês.

    **4. Palavras-chave (em português)**
    - De 3 a 5 palavras separadas por ponto e vírgula (;).

    **5. Introdução**
    - Apresente o tema, justificativa, problema e objetivo da pesquisa.
    - Mínimo de 200 palavras.

    **6. Revisão de Literatura**
    - Discorra sobre conceitos teóricos importantes sobre o tema.
    - Utilize ao menos 2 citações no estilo ABNT: (SOBRENOME, ano, p.xx).
    - Mínimo de 300 palavras.

    **7. Metodologia**
    - Descreva os métodos e procedimentos adotados para desenvolver o trabalho.
    - Pode incluir abordagem qualitativa/quantitativa, revisão bibliográfica, etc.
    - Mínimo de 200 palavras.

    **8. Resultados e Discussão**
    - Apresente os principais resultados esperados ou obtidos.
    - Relacione com a literatura citada.
    - Mínimo de 300 palavras.

    **9. Conclusão**
    - Retome os objetivos, destaque as contribuições e proponha trabalhos futuros.
    - Mínimo de 150 palavras.

    **10. Referências**
    - Liste pelo menos **3 referências no formato ABNT.**
    - Exemplo: SOBRENOME, Nome. *Título do Livro ou Artigo*. Local: Editora, Ano.

    **Instruções finais:**
    - Escreva de forma acadêmica, coesa e formal.
    - Use normas da ABNT para estruturação e citações.
    - Retorne o texto organizado por seções, com títulos bem definidos.
    """

    resposta = modelo.generate_content(prompt)
    return resposta.text

# Função para salvar o texto no PDF com formatação correta
def salvar_em_pdf(titulo, texto):
    nome_arquivo = titulo.replace(" ", "_") + ".pdf"
    doc = SimpleDocTemplate(nome_arquivo, pagesize=A4,
                            rightMargin=2*cm, leftMargin=2*cm,
                            topMargin=2*cm, bottomMargin=2*cm)

    # Estilos
    styles = getSampleStyleSheet()
    styles.add(ParagraphStyle(name='Titulo', fontSize=16, spaceAfter=12, alignment=1))
    styles.add(ParagraphStyle(name='Secao', fontSize=14, spaceAfter=8, spaceBefore=12, leading=18))
    styles.add(ParagraphStyle(name='Texto', fontSize=12, leading=16, alignment=TA_JUSTIFY))

    elementos = []

    # Título do trabalho
    elementos.append(Paragraph(titulo, styles['Titulo']))

    # Regex para detectar seções
    padroes = {
        "Resumo": re.compile(r"^(?:\*+)?\s*\d{0,1}\.?\s*resumo[:\*]*\s*$", re.IGNORECASE),
        "Abstract": re.compile(r"^(?:\*+)?\s*\d{0,1}\.?\s*abstract[:\*]*\s*$", re.IGNORECASE),
        "Palavras-chave": re.compile(r"^(?:\*+)?\s*\d{0,1}\.?\s*palavras-chave[:\*]*\s*$", re.IGNORECASE),
        "Introdução": re.compile(r"^(?:\*+)?\s*\d{0,1}\.?\s*introdução[:\*]*\s*$", re.IGNORECASE),
        "Revisão de Literatura": re.compile(r"^(?:\*+)?\s*\d{0,1}\.?\s*revisão de literatura[:\*]*\s*$", re.IGNORECASE),
        "Metodologia": re.compile(r"^(?:\*+)?\s*\d{0,1}\.?\s*metodologia[:\*]*\s*$", re.IGNORECASE),
        "Resultados e Discussão": re.compile(r"^(?:\*+)?\s*\d{0,1}\.?\s*resultados e discussão[:\*]*\s*$", re.IGNORECASE),
        "Conclusão": re.compile(r"^(?:\*+)?\s*\d{0,1}\.?\s*conclusão[:\*]*\s*$", re.IGNORECASE),
        "Referências": re.compile(r"^(?:\*+)?\s*\d{0,1}\.?\s*refer[eê]ncias[:\*]*\s*$", re.IGNORECASE)
    }

    conteudo = {}
    atual = None

    # Remover os três asteriscos e processar as seções
    for linha in texto.splitlines():
        linha_limpa = remover_asteriscos(linha.strip())
        if not linha_limpa:
            continue

        mudou_secao = False
        for nome_secao, padrao in padroes.items():
            if padrao.match(linha_limpa):
                atual = nome_secao
                conteudo[atual] = []
                mudou_secao = True
                break

        if not mudou_secao and atual:
            conteudo[atual].append(linha_limpa)

    # Adicionar conteúdo por seção
    for secao in ["Resumo", "Abstract", "Palavras-chave", "Introdução", "Revisão de Literatura", "Metodologia", "Resultados e Discussão", "Conclusão", "Referências"]:
        elementos.append(Paragraph(secao.upper(), styles['Secao']))
        if secao in conteudo and conteudo[secao]:
            for paragrafo in conteudo[secao]:
                elementos.append(Paragraph(paragrafo, styles['Texto']))
                elementos.append(Spacer(1, 0.4*cm))
        else:
            elementos.append(Paragraph("Conteúdo não disponível.", styles['Texto']))

    doc.build(elementos)
    print(f"\n✅ PDF salvo como: {nome_arquivo}")

# Função para salvar no DOCX com formatação correta
def salvar_em_docx(titulo, texto):
    doc = Document()

    # Título principal
    titulo_paragrafo = doc.add_heading(titulo, level=1)
    titulo_paragrafo.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Regex para detectar seções
    padroes = {
        "Resumo": re.compile(r"^(?:\*+)?\s*\d{0,1}\.?\s*resumo[:\*]*\s*$", re.IGNORECASE),
        "Abstract": re.compile(r"^(?:\*+)?\s*\d{0,1}\.?\s*abstract[:\*]*\s*$", re.IGNORECASE),
        "Palavras-chave": re.compile(r"^(?:\*+)?\s*\d{0,1}\.?\s*palavras-chave[:\*]*\s*$", re.IGNORECASE),
        "Introdução": re.compile(r"^(?:\*+)?\s*\d{0,1}\.?\s*introdução[:\*]*\s*$", re.IGNORECASE),
        "Revisão de Literatura": re.compile(r"^(?:\*+)?\s*\d{0,1}\.?\s*revisão de literatura[:\*]*\s*$", re.IGNORECASE),
        "Metodologia": re.compile(r"^(?:\*+)?\s*\d{0,1}\.?\s*metodologia[:\*]*\s*$", re.IGNORECASE),
        "Resultados e Discussão": re.compile(r"^(?:\*+)?\s*\d{0,1}\.?\s*resultados e discussão[:\*]*\s*$", re.IGNORECASE),
        "Conclusão": re.compile(r"^(?:\*+)?\s*\d{0,1}\.?\s*conclusão[:\*]*\s*$", re.IGNORECASE),
        "Referências": re.compile(r"^(?:\*+)?\s*\d{0,1}\.?\s*refer[eê]ncias[:\*]*\s*$", re.IGNORECASE)
    }

    conteudo = {}
    atual = None

    # Remover os três asteriscos e processar as seções
    for linha in texto.splitlines():
        linha_limpa = remover_asteriscos(linha.strip())
        if not linha_limpa:
            continue

        mudou_secao = False
        for nome_secao, padrao in padroes.items():
            if padrao.match(linha_limpa):
                atual = nome_secao
                conteudo[atual] = []
                mudou_secao = True
                break

        if not mudou_secao and atual:
            conteudo[atual].append(linha_limpa)

    # Adicionar conteúdo por seção
    for secao in ["Resumo", "Abstract", "Palavras-chave", "Introdução", "Revisão de Literatura", "Metodologia", "Resultados e Discussão", "Conclusão", "Referências"]:
        doc.add_heading(secao.upper(), level=2) 
        if secao in conteudo and conteudo[secao]:
            for paragrafo in conteudo[secao]:
                p = doc.add_paragraph(paragrafo)
                run = p.runs[0]
                run.font.name = "Times New Roman"
                run.font.size = Pt(12)
                run.font.color.rgb = RGBColor(0, 0, 0)
        else:
            doc.add_paragraph("Conteúdo não disponível.")

    # Aplicar formatação geral
    formatar_paragrafos(doc)

    nome_arquivo = titulo.replace(" ", "_") + ".docx"
    doc.save(nome_arquivo)
    print(f"\n✅ Documento salvo como: {nome_arquivo}")

# Coletar dados do usuário
def coletar_dados_usuario():
    while True:
        titulo = input("Digite o título do trabalho: ").strip()
        if titulo:
            titulo = titulo.title()  
            break
        print("⚠️ O título não pode estar vazio. Tente novamente.\n")

    while True:
        tema = input("Digite o tema do trabalho: ").strip()
        if tema:
            break
        print("⚠️ O tema não pode estar vazio. Tente novamente.\n")

    return titulo, tema

# Execução
titulo, tema = coletar_dados_usuario()
trabalho = gerar_artigo_abnt(titulo, tema)

print("\n========= ARTIGO GERADO =========\n")
print(trabalho)
print("\n===============================\n")

formato = input("Deseja gerar em DOCX, PDF ou ambos? ").lower()

if "docx" in formato or "ambos" in formato:
    salvar_em_docx(titulo, trabalho)

if "pdf" in formato or "ambos" in formato:
    salvar_em_pdf(titulo, trabalho)
